from __future__ import annotations

from datetime import datetime
from pathlib import Path

from core.transcribe import TranscriptSegment


def write_docx(
    segments: list[TranscriptSegment],
    output_path: str,
    title: str = "NoisyWhisper Transcript",
) -> Path:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. Install dependencies first."
        ) from exc

    target = Path(output_path).expanduser().resolve()
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(datetime.now().strftime("Generated: %Y-%m-%d %H:%M:%S"))

    for seg in segments:
        text = seg.text.strip()
        if not text:
            continue
        start_s = seg.start_ms // 1000
        hh = start_s // 3600
        mm = (start_s % 3600) // 60
        ss = start_s % 60
        doc.add_paragraph(f"[{hh:02d}:{mm:02d}:{ss:02d}] {text}")

    doc.save(str(target))
    return target
